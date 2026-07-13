from __future__ import annotations

import html
import re
import subprocess
import tempfile
import zipfile
from decimal import Decimal, ROUND_HALF_UP
from io import BytesIO
from pathlib import Path
from xml.etree import ElementTree as ET

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# OOXML namespaces (proposal preview reads raw XML when python-docx body text is sparse).
_W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _w_tag(local: str) -> str:
    return "{%s}%s" % (_W_NS, local)

_D0 = Decimal("0")
_CENT = Decimal("0.01")

# Canonical Word tokens ({{TOKEN}}) after normalization + value fill.
_DOCX_CANONICAL_TOKENS: tuple[str, ...] = (
    "JOB_NAME",
    "QUOTE_NUMBER",
    "CUSTOMER_NAME",
    "CONTACT_NAME",
    "PROPOSAL_AMOUNT",
    "SCOPE_OF_WORK",
    "CUSTOMER_RESPONSIBILITIES",
    "PREPARED_BY",
    "DATE",
    "PREPARED_BY_PHONE",
    "ESTIMATE_DATE",
    "VALID_THROUGH",
    "EXPIRATION_DATE",
)

ESTIMATE_PROPOSAL_TEMPLATE_FILENAME = "estimate_template_autofill_logo_updated.docx"

# Default phone in quote footer when ``PREPARED_BY_PHONE`` is blank (matches IPS letterhead).
IPS_DEFAULT_QUOTE_FOOTER_PHONE = "(337) 577-3944"

# Standard company logo for proposals (first existing file wins). Same folder as the app ``assets``.
# Optional Word placeholder: ``{{COMPANY_LOGO}}`` in header/body/footer — replaced with this image when present.
COMPANY_LOGO_CANDIDATE_FILENAMES: tuple[str, ...] = (
    "company_logo.png",
    "company_logo.jpg",
    "company_logo.jpeg",
    "ips_logo_wide.png",
    "IPS LOGO WIDE.png",
)

_PROPOSAL_LOGO_PLACEHOLDER_RE = re.compile(r"\{\{\s*COMPANY\s*LOGO\s*\}\}", re.IGNORECASE)

# Shown when LibreOffice / Word PDF conversion is not available (UI may prepend context).
PROPOSAL_PDF_UNAVAILABLE_MSG = (
    "PDF export is not available in this environment (Word could not be converted to PDF). "
    "You can still use **Download Word Proposal** and save or print to PDF locally. "
    "To enable server PDFs, install **LibreOffice** so `soffice` is on PATH, or on Windows use **Microsoft Word** with the **docx2pdf** package."
)


def _dec(v) -> Decimal:
    if v is None or v == "":
        return _D0
    if isinstance(v, Decimal):
        return v
    return Decimal(str(v))


def _q2(v) -> Decimal:
    return _dec(v).quantize(_CENT, rounding=ROUND_HALF_UP)


def _money(v) -> str:
    """Format as $X,XXX.XX (always two decimals)."""
    return f"${_q2(v):,.2f}"


def _s(v) -> str:
    """Safe string for placeholders; never None."""
    if v is None:
        return ""
    return str(v).strip() if isinstance(v, str) else str(v)


def _s_contact(v) -> str:
    """Contact / Attn line: never show literal 'None' or placeholder noise."""
    s = _s(v)
    if not s:
        return ""
    low = s.lower()
    if low in ("none", "null"):
        return ""
    return s


def _pdf_prepared_by(est: dict) -> str:
    try:
        from app.auth import current_profile

        prof = current_profile()
        name = _s(prof.get("full_name"))
        if name:
            return name
        email = _s(prof.get("email"))
        if email:
            return email
    except Exception:
        pass
    return _s(est.get("prepared_by_name") or est.get("prepared_by"))


def _proposal_prepared_by_display(est: dict) -> str:
    n = _s(est.get("prepared_by_name"))
    if n:
        return n
    return _pdf_prepared_by(est)


def _paragraph_heading_level(paragraph) -> int:
    """Map Word paragraph style to a rough heading level for HTML preview."""
    try:
        name = (paragraph.style.name or "").strip().lower()
    except Exception:
        name = ""
    if name in ("title",):
        return 1
    if name in ("subtitle",):
        return 2
    if name.startswith("heading"):
        parts = name.split()
        if len(parts) >= 2 and parts[1].isdigit():
            return min(max(int(parts[1]), 1), 4)
        return 2
    return 0


def _iter_document_body_blocks(document: Document):
    """Yield (``'p'``, paragraph) or (``'t'``, table) in document body order."""
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    body = document.element.body
    for el in body.iterchildren():
        if isinstance(el, CT_P):
            yield "p", Paragraph(el, document)
        elif isinstance(el, CT_Tbl):
            yield "t", Table(el, document)


def _cell_plain_text(cell) -> str:
    parts: list[str] = []
    for p in cell.paragraphs:
        t = (p.text or "").strip()
        if t:
            parts.append(t)
    return " ".join(parts).strip()


def _table_to_preview_html(table) -> str:
    rows_out: list[list[str]] = []
    for row in table.rows:
        row_cells = [_cell_plain_text(c) for c in row.cells]
        if any(row_cells):
            rows_out.append(row_cells)
    if not rows_out:
        return ""
    ncol = max(len(r) for r in rows_out)
    trs: list[str] = []
    for r in rows_out:
        padded = r + [""] * (ncol - len(r))
        tds = "".join(
            f"<td class='ips-proposal-docx-td'>{html.escape(c)}</td>"
            for c in padded
        )
        trs.append(f"<tr>{tds}</tr>")
    return (
        "<table class='ips-proposal-docx-table' style='width:100%;border-collapse:collapse;margin:0.55em 0;"
        "font-size:0.95em;'>"
        "<tbody>" + "".join(trs) + "</tbody></table>"
    )


def _first_header_preview_html(document: Document) -> str:
    """Optional letterhead lines from the first section header (text only)."""
    try:
        hdr = document.sections[0].header
    except Exception:
        return ""
    lines: list[str] = []
    for p in hdr.paragraphs:
        t = (p.text or "").strip()
        if t:
            lines.append(html.escape(t))
    if not lines:
        return ""
    inner = "<br/>".join(lines)
    return f'<div class="ips-proposal-letterhead">{inner}</div>'


def _paragraph_text_from_ooxml_w_p(p_el: ET.Element) -> str:
    """Plain text from a ``w:p`` element (handles split runs, tabs, breaks, simple revisions)."""
    parts: list[str] = []
    for node in p_el.iter():
        tag = node.tag
        if tag in (_w_tag("t"), _w_tag("delText")):
            if node.text:
                parts.append(node.text)
        elif tag == _w_tag("tab"):
            parts.append("\t")
        elif tag in (_w_tag("br"), _w_tag("cr")):
            parts.append("\n")
    return "".join(parts).replace("\x07", "").strip()


def _table_html_from_ooxml_w_tbl(tbl_el: ET.Element) -> str:
    rows_out: list[list[str]] = []
    for tr in tbl_el.findall(_w_tag("tr")):
        row_cells: list[str] = []
        for tc in tr.findall(_w_tag("tc")):
            cell_parts: list[str] = []
            for t in tc.findall(".//" + _w_tag("t")):
                if t.text:
                    cell_parts.append(t.text)
            row_cells.append("".join(cell_parts).strip())
        if any(c for c in row_cells):
            rows_out.append(row_cells)
    if not rows_out:
        return ""
    ncol = max(len(r) for r in rows_out)
    trs: list[str] = []
    for r in rows_out:
        padded = r + [""] * (ncol - len(r))
        tds = "".join(
            "<td class='ips-proposal-docx-td'>"
            f"{html.escape(c)}</td>"
            for c in padded
        )
        trs.append(f"<tr>{tds}</tr>")
    return (
        "<table class='ips-proposal-docx-table' style='width:100%;border-collapse:collapse;margin:0.55em 0;"
        "font-size:0.95em;'>"
        "<tbody>" + "".join(trs) + "</tbody></table>"
    )


def _ooxml_headers_footers_preview_html(zf: zipfile.ZipFile) -> str:
    """Letterhead / footer text from ``word/header*.xml`` and ``word/footer*.xml``."""
    lines: list[str] = []
    for name in sorted(zf.namelist()):
        if not (name.startswith("word/header") or name.startswith("word/footer")):
            continue
        if not name.endswith(".xml"):
            continue
        try:
            root = ET.fromstring(zf.read(name))
        except Exception:
            continue
        for p in root.findall(_w_tag("p")):
            t = _paragraph_text_from_ooxml_w_p(p)
            if t:
                lines.append(html.escape(t))
    if not lines:
        return ""
    inner = "<br/>".join(lines[:48])
    return f'<div class="ips-proposal-letterhead">{inner}</div>'


def _ooxml_local_tag(el: ET.Element) -> str:
    t = el.tag
    return t.split("}", 1)[-1] if "}" in t else t


def _iter_ooxml_body_blocks(body: ET.Element):
    """
    Yield ``w:p`` and ``w:tbl`` elements from ``w:body`` in order, including inside ``w:sdt`` /
    ``w:sdtContent`` (common for protected regions and content controls in proposal templates).
    """
    for child in list(body):
        tag = _ooxml_local_tag(child)
        if tag == "sectPr":
            return
        if tag in ("p", "tbl"):
            yield child
        elif tag == "sdt":
            sdtc = child.find(_w_tag("sdtContent"))
            if sdtc is not None:
                yield from _iter_ooxml_body_blocks(sdtc)


def _ooxml_drawingml_supplement_chunks(doc_root: ET.Element, *, min_chars: int = 12) -> list[str]:
    """Extra visible text from DrawingML ``a:t`` runs (text boxes / shapes) when body text is sparse."""
    a_ns = "http://schemas.openxmlformats.org/drawingml/2006/main"
    seen: set[str] = set()
    out: list[str] = []
    for t in doc_root.findall(".//{%s}t" % a_ns):
        tx = " ".join(str(t.text or "").split()).strip()
        if len(tx) < min_chars or tx in seen:
            continue
        seen.add(tx)
        esc = html.escape(tx).replace("\n", "<br/>")
        out.append(f'<p style="margin:0.22em 0;line-height:1.52;opacity:0.95;">{esc}</p>')
    return out[:24]


def _preview_inner_html_from_docx_ooxml(docx_bytes: bytes) -> str:
    """
    Build inner HTML from the raw OOXML package (``word/document.xml``).

    Catches paragraph text that lives only in split ``w:r``/``w:t`` nodes, which ``python-docx``
    sometimes surfaces as empty ``paragraph.text`` in heavily styled templates.
    Walks ``w:sdt`` / ``w:sdtContent`` so content-control templates still preview.
    """
    chunks: list[str] = []
    doc_root: ET.Element | None = None
    try:
        with zipfile.ZipFile(BytesIO(docx_bytes), "r") as zf:
            if "word/document.xml" not in zf.namelist():
                return ""
            hf = _ooxml_headers_footers_preview_html(zf)
            if hf:
                chunks.append(hf)
            doc_root = ET.fromstring(zf.read("word/document.xml"))
    except Exception:
        return ""

    if doc_root is None:
        return ""
    body = doc_root.find(_w_tag("body"))
    if body is None:
        return ""
    for child in _iter_ooxml_body_blocks(body):
        tag = _ooxml_local_tag(child)
        if tag == "p":
            text = _paragraph_text_from_ooxml_w_p(child)
            if not text:
                continue
            esc = html.escape(text).replace("\n", "<br/>")
            chunks.append(f'<p style="margin:0.22em 0;line-height:1.52;">{esc}</p>')
        elif tag == "tbl":
            th = _table_html_from_ooxml_w_tbl(child)
            if th:
                chunks.append(th)
    inner = "\n".join(chunks)
    if _preview_plain_text_len(inner) < 40 and doc_root is not None:
        chunks.extend(_ooxml_drawingml_supplement_chunks(doc_root))
        inner = "\n".join(chunks)
    return inner


def _preview_inner_html_from_python_docx(docx_bytes: bytes) -> str:
    """Body preview via ``python-docx`` (headings + tables) — secondary to OOXML for coverage."""
    doc = Document(BytesIO(docx_bytes))
    chunks: list[str] = []
    hf = _first_header_preview_html(doc)
    if hf:
        chunks.append(hf)
    for kind, block in _iter_document_body_blocks(doc):
        if kind == "p":
            text = (block.text or "").replace("\x07", "").strip()
            if not text:
                continue
            esc = html.escape(text).replace("\n", "<br/>")
            lev = _paragraph_heading_level(block)
            if lev == 1:
                chunks.append(f'<h3 style="margin:0.55em 0 0.25em 0;font-weight:650;">{esc}</h3>')
            elif lev == 2:
                chunks.append(f'<h4 style="margin:0.45em 0 0.2em 0;font-weight:600;">{esc}</h4>')
            elif lev >= 3:
                chunks.append(f'<h5 style="margin:0.35em 0 0.15em 0;font-weight:600;">{esc}</h5>')
            else:
                chunks.append(f'<p style="margin:0.22em 0;line-height:1.52;">{esc}</p>')
        else:
            th = _table_to_preview_html(block)
            if th:
                chunks.append(th)
    return "\n".join(chunks)


def _preview_plain_text_len(inner_html: str) -> int:
    plain = re.sub(r"<[^>]+>", " ", inner_html or "")
    plain = re.sub(r"\s+", " ", plain).strip()
    return len(plain)


def _wrap_docx_preview(inner: str) -> str:
    inner = (inner or "").strip()
    if not inner:
        return ""
    return (
        '<div class="ips-proposal-docx-body" style="max-height:min(720px,78vh);overflow:auto;'
        "margin:0;padding:0;font-family:Georgia,\"Times New Roman\",serif;\">"
        f"{inner}</div>"
    )


def _best_preview_inner_html_from_docx_bytes(docx_bytes: bytes) -> str:
    """Pick the richest HTML extraction for the same filled DOCX bytes used for download."""
    xml_inner = _preview_inner_html_from_docx_ooxml(docx_bytes)
    pc_inner = _preview_inner_html_from_python_docx(docx_bytes)
    if _preview_plain_text_len(xml_inner) >= _preview_plain_text_len(pc_inner):
        return xml_inner.strip()
    return pc_inner.strip()


def proposal_values_preview_html(vals: dict[str, str]) -> str:
    """Field-based preview (same placeholder map as the DOCX / :func:`proposal_values`). All rows always render."""
    # Same keys as ``proposal_values`` — always show a row (empty → em dash) so preview matches Word placeholders.
    row_order: tuple[tuple[str, str], ...] = (
        ("JOB_NAME", "Project title (estimate description)"),
        ("QUOTE_NUMBER", "Quote #"),
        ("CUSTOMER_NAME", "Customer"),
        ("CONTACT_NAME", "Contact"),
        ("PROPOSAL_AMOUNT", "Proposal amount"),
        ("SCOPE_OF_WORK", "Scope of work"),
        ("CUSTOMER_RESPONSIBILITIES", "Customer responsibilities"),
        ("PREPARED_BY", "Prepared by"),
        ("DATE", "Date"),
        ("ESTIMATE_DATE", "Estimate date"),
        ("VALID_THROUGH", "Valid through"),
        ("EXPIRATION_DATE", "Expiration date"),
        ("PREPARED_BY_PHONE", "Phone"),
    )
    trs: list[str] = []
    for key, label in row_order:
        raw = _s(vals.get(key, ""))
        lab_esc = html.escape(label)
        val_esc = html.escape(raw) if raw else "—"
        trs.append(
            "<tr>"
            f"<td class='ips-proposal-fallback-label'>{lab_esc}</td>"
            f"<td class='ips-proposal-fallback-val'>{val_esc}</td>"
            "</tr>"
        )
    body = "".join(trs)
    return (
        '<div class="ips-proposal-fields-wrap">'
        f"<table class='ips-proposal-fallback-table'><tbody>{body}</tbody></table></div>"
    )


def proposal_preview_page_html(view_model=None):
    """
    HTML preview shell aligned with :mod:`app.estimate.proposal_document_layout`
    (same structured layout as Word exports — not OOXML extraction).

    Pass ``None`` when no view model can be built (shows an instructional placeholder page).
    """
    from app.estimate.proposal_document_layout import render_proposal_preview_page_html
    return render_proposal_preview_page_html(view_model)


def proposal_preview_html(view_model=None, *, fallback_vals: dict[str, str] | None = None) -> str:
    """Backward-compatible wrapper; ``fallback_vals`` is ignored."""
    _ = fallback_vals
    return proposal_preview_page_html(view_model)


def proposal_values(
    est: dict,
    totals: dict,
    customer_name: str = "",
    job_name: str = "",
    *,
    contact_name: str = "",
    prepared_by_phone: str = "",
) -> dict[str, str]:
    """
    Placeholder values for the estimate proposal DOCX (and PDF derived from it).

    Built from :func:`app.estimate.proposal_document_layout.build_proposal_view_model` so ordering matches preview HTML.

    Money: $X,XXX.XX. Date: full US month name. Missing fields -> empty string (or sensible default for title).
    """
    from app.estimate.proposal_document_layout import (
        build_proposal_view_model,
        proposal_placeholder_map_from_view_model,
    )
    vm = build_proposal_view_model(
        est,
        totals,
        customer_name=customer_name,
        job_name=job_name,
        contact_name=contact_name,
        prepared_by_phone=prepared_by_phone,
    )
    return proposal_placeholder_map_from_view_model(vm)


def _proposal_assets_dir() -> Path:
    return Path(__file__).resolve().parents[1] / "assets"


def _default_estimate_proposal_template_path() -> Path:
    """IPS APP/assets/estimate_template_autofill_logo_updated.docx"""
    return _proposal_assets_dir() / ESTIMATE_PROPOSAL_TEMPLATE_FILENAME


def _load_standard_proposal_template_bytes() -> bytes:
    """
    Load the single packaged proposal template (no overrides):
    ``assets/estimate_template_autofill_logo_updated.docx``.
    """
    default = _default_estimate_proposal_template_path()
    if default.is_file():
        return default.read_bytes()
    raise FileNotFoundError(
        f"Proposal template not found. Add {ESTIMATE_PROPOSAL_TEMPLATE_FILENAME} under the project "
        f"**assets** folder ({default.parent})."
    )


def _resolve_standard_company_logo_path() -> Path | None:
    """First matching file under ``assets/``; ``None`` if no standard logo is shipped."""
    root = _proposal_assets_dir()
    for name in COMPANY_LOGO_CANDIDATE_FILENAMES:
        p = root / name
        if p.is_file():
            return p
    return None


def _clear_paragraph_runs(paragraph) -> None:
    for r in list(paragraph.runs):
        paragraph._element.remove(r._element)


def _paragraph_full_text(paragraph) -> str:
    if paragraph.runs:
        return "".join(r.text or "" for r in paragraph.runs)
    return paragraph.text or ""


def _replace_company_logo_placeholders_in_container(container, logo_path: Path, width) -> int:
    """Replace ``{{COMPANY_LOGO}}`` paragraphs with an inline image. Returns count replaced."""
    replaced = 0
    for p in container.paragraphs:
        if not _PROPOSAL_LOGO_PLACEHOLDER_RE.search(_paragraph_full_text(p)):
            continue
        _clear_paragraph_runs(p)
        run = p.add_run()
        run.add_picture(str(logo_path), width=width)
        replaced += 1
    for tbl in getattr(container, "tables", []) or []:
        for row in tbl.rows:
            for cell in row.cells:
                replaced += _replace_company_logo_placeholders_in_container(cell, logo_path, width)
    return replaced


def _replace_company_logo_placeholders_in_document(doc: Document, logo_path: Path, width) -> int:
    n = 0
    for p in doc.paragraphs:
        if not _PROPOSAL_LOGO_PLACEHOLDER_RE.search(_paragraph_full_text(p)):
            continue
        _clear_paragraph_runs(p)
        p.add_run().add_picture(str(logo_path), width=width)
        n += 1
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                n += _replace_company_logo_placeholders_in_container(cell, logo_path, width)
    for section in doc.sections:
        for part in (section.header, section.footer):
            n += _replace_company_logo_placeholders_in_container(part, logo_path, width)
    return n


def _prepend_company_logo_primary_header(doc: Document, logo_path: Path, width) -> None:
    """Insert a logo row above the first header paragraph (section 0 only) for consistent letterhead."""
    if not doc.sections:
        return
    hdr = doc.sections[0].header
    if not hdr.paragraphs:
        hdr.add_paragraph()
    anchor = hdr.paragraphs[0]
    logo_p = anchor.insert_paragraph_before("")
    logo_p.add_run().add_picture(str(logo_path), width=width)


def _apply_standard_proposal_branding(doc: Document) -> None:
    """
    Apply the standard company logo to the generated document.

    - If ``{{COMPANY_LOGO}}`` appears in the template, those blocks become the canonical logo image.
    - Otherwise, when a standard logo file exists under ``assets/``, it is prepended to the primary header.
    """
    logo_path = _resolve_standard_company_logo_path()
    if not logo_path:
        return
    width = Inches(7.5)  # Full content width (8.5in page − 0.5in margins); sync with HTML ``.ips-ph-logo-img``.
    replaced = _replace_company_logo_placeholders_in_document(doc, logo_path, width)
    if replaced == 0:
        try:
            _prepend_company_logo_primary_header(doc, logo_path, width)
        except Exception:
            pass


def _normalize_placeholder_tokens(text: str) -> str:
    if not text:
        return text
    s = text
    explicit: list[tuple[str, str]] = [
        (r"\{\{\s*Customer\s+Name\s*\}\}", "{{CUSTOMER_NAME}}"),
        (r"\{\{\s*Scope\s+of\s+Work\s*\}\}", "{{SCOPE_OF_WORK}}"),
        (r"\{\{\s*Customer\s+Responsibilities\s*\}\}", "{{CUSTOMER_RESPONSIBILITIES}}"),
        (r"\{\{\s*Job\s*Name\s*\}\}", "{{JOB_NAME}}"),
        (r"\{\{\s*Quote\s*Number\s*\}\}", "{{QUOTE_NUMBER}}"),
        (r"\{\{\s*Contact\s*Name\s*\}\}", "{{CONTACT_NAME}}"),
        (r"\{\{\s*Proposal\s*Amount\s*\}\}", "{{PROPOSAL_AMOUNT}}"),
        (r"\{\{\s*Prepared\s*By\s*\}\}", "{{PREPARED_BY}}"),
        (r"\{\{\s*Prepared\s*By\s*Phone\s*\}\}", "{{PREPARED_BY_PHONE}}"),
        (r"\{\{\s*Prepared\s*By\s*Phone\s*Number\s*\}\}", "{{PREPARED_BY_PHONE}}"),
        (r"\{\{\s*Date\s*\}\}", "{{DATE}}"),
        (r"\{\{\s*Estimate\s*Date\s*\}\}", "{{ESTIMATE_DATE}}"),
        (r"\{\{\s*Valid\s*Through\s*\}\}", "{{VALID_THROUGH}}"),
        (r"\{\{\s*Expiration\s*Date\s*\}\}", "{{EXPIRATION_DATE}}"),
        (r"\{\{\s*Company\s+Logo\s*\}\}", "{{COMPANY_LOGO}}"),
    ]
    for pat, repl in explicit:
        s = re.sub(pat, repl, s, flags=re.IGNORECASE)

    for tok in _DOCX_CANONICAL_TOKENS:
        s = re.sub(
            r"\{\{\s*" + re.escape(tok) + r"\s*\}\}",
            "{{" + tok + "}}",
            s,
            flags=re.IGNORECASE,
        )
    return s


def _docx_placeholder_replacements(vals: dict[str, str]) -> dict[str, str]:
    repl: dict[str, str] = {}
    for k in _DOCX_CANONICAL_TOKENS:
        v = _s(vals.get(k, ""))
        if k == "PREPARED_BY_PHONE" and not v:
            v = IPS_DEFAULT_QUOTE_FOOTER_PHONE
        repl["{{" + k + "}}"] = v
    return repl


def _set_paragraph_text_merged(p, new_text: str) -> None:
    """
    Set paragraph text after Word split content across runs.

    Merge is logical only: remove **all** runs, then add **one** run with the final string so
    placeholders that span runs are replaced reliably and the saved OOXML stays consistent.
    """
    tx = new_text if new_text is not None else ""
    _clear_paragraph_runs(p)
    p.add_run(tx)


def _walk_table_paragraphs(table):
    """Yield every paragraph inside a table, including nested tables."""
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                yield p
            for tbl in cell.tables:
                yield from _walk_table_paragraphs(tbl)


def _iter_all_paragraphs(doc: Document):
    """Body, tables, headers, and footers — every paragraph that can hold {{PLACEHOLDERS}}."""
    for p in doc.paragraphs:
        yield p
    for tbl in doc.tables:
        yield from _walk_table_paragraphs(tbl)
    for section in doc.sections:
        for part in (section.header, section.footer):
            for p in part.paragraphs:
                yield p
            for tbl in getattr(part, "tables", []) or []:
                yield from _walk_table_paragraphs(tbl)


def _replace_placeholders_in_text(text: str, replacements: dict[str, str]) -> str:
    """
    Substitute all canonical ``{{TOKEN}}`` occurrences with values.

    Uses exact keys first, then case-insensitive regex so spacing variants inside braces still match.
    """
    new = text
    for tok in _DOCX_CANONICAL_TOKENS:
        k = "{{" + tok + "}}"
        v = replacements.get(k, "")
        if k in new:
            new = new.replace(k, v)
        pat = re.compile(r"\{\{\s*" + re.escape(tok) + r"\s*\}\}", re.IGNORECASE)
        new = pat.sub(lambda _m, val=v: val, new)
    return new


def _apply_to_paragraph_text(doc: Document, fn) -> None:
    """Apply ``fn(full_paragraph_text) -> new_text`` to every paragraph; merged run write on change."""

    def apply_paragraph(p) -> None:
        text = _paragraph_full_text(p)
        if not text:
            return
        new = fn(text)
        if new != text:
            _set_paragraph_text_merged(p, new)

    for p in _iter_all_paragraphs(doc):
        apply_paragraph(p)


def _normalize_docx_placeholders(doc: Document) -> None:
    _apply_to_paragraph_text(doc, _normalize_placeholder_tokens)


def _replace_placeholders_doc(doc: Document, replacements: dict[str, str]) -> None:
    def replace(text: str) -> str:
        return _replace_placeholders_in_text(text, replacements)

    _apply_to_paragraph_text(doc, replace)


def _remove_customer_location_paragraphs(doc: Document) -> None:
    """Strip legacy ``Customer Location: …`` blocks so quotes never show site lines."""
    kill: list = []
    for p in _iter_all_paragraphs(doc):
        t = (_paragraph_full_text(p) or "").strip()
        if re.match(r"^Customer Location\s*:", t, flags=re.IGNORECASE):
            kill.append(p._element)
    for el in kill:
        parent = el.getparent()
        if parent is not None:
            parent.remove(el)


def _p_ensure_pPr(p_el) -> object:
    pPr = p_el.find(_w_tag("pPr"))
    if pPr is None:
        pPr = OxmlElement("w:pPr")
        p_el.insert(0, pPr)
    return pPr


def _paragraph_set_shading(paragraph, fill_hex: str) -> None:
    """Paragraph background (OOXML ``w:shd``); ``fill_hex`` like ``0B4F8A``."""
    pPr = _p_ensure_pPr(paragraph._element)
    for el in list(pPr):
        if el.tag == _w_tag("shd"):
            pPr.remove(el)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), fill_hex)
    pPr.append(shd)


def _paragraph_clear_shading(paragraph) -> None:
    pPr = paragraph._element.find(_w_tag("pPr"))
    if pPr is None:
        return
    for el in list(pPr):
        if el.tag == _w_tag("shd"):
            pPr.remove(el)


def _paragraph_set_bottom_border_double(paragraph) -> None:
    pPr = _p_ensure_pPr(paragraph._element)
    pBdr = pPr.find(_w_tag("pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    for el in list(pBdr):
        if el.tag == _w_tag("bottom"):
            pBdr.remove(el)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "double")
    bottom.set(qn("w:sz"), "12")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)


def _paragraph_set_top_border_double(paragraph) -> None:
    pPr = _p_ensure_pPr(paragraph._element)
    pBdr = pPr.find(_w_tag("pBdr"))
    if pBdr is None:
        pBdr = OxmlElement("w:pBdr")
        pPr.append(pBdr)
    for el in list(pBdr):
        if el.tag == _w_tag("top"):
            pBdr.remove(el)
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "double")
    top.set(qn("w:sz"), "12")
    top.set(qn("w:space"), "6")
    top.set(qn("w:color"), "000000")
    pBdr.append(top)


def _runs_set_font(paragraph, *, name: str, size_pt: float, bold: bool = False, italic: bool = False) -> None:
    for r in paragraph.runs:
        r.font.name = name
        r.font.size = Pt(size_pt)
        r.font.bold = bold
        r.font.italic = italic
        r.font.color.rgb = RGBColor(0, 0, 0)


def _runs_set_banner_white(paragraph, *, size_pt: float) -> None:
    for r in paragraph.runs:
        r.font.name = "Arial"
        r.font.size = Pt(size_pt)
        r.font.bold = True
        r.font.italic = True
        r.font.color.rgb = RGBColor(255, 255, 255)


def _is_quote_title_line(t: str) -> bool:
    s = (t or "").strip()
    if "Quote #:" in s:
        return False
    return bool(re.search(r"[\u2013\u2014\-]\s*Quote\s*$", s))


def _format_intro_solutions_underline(paragraph) -> None:
    full = (_paragraph_full_text(paragraph) or "").strip()
    if "Solutions" not in full:
        _runs_set_font(paragraph, name="Times New Roman", size_pt=11.0)
        return
    before, _, after = full.partition("Solutions")
    _clear_paragraph_runs(paragraph)
    r0 = paragraph.add_run(before)
    r1 = paragraph.add_run("Solutions")
    r2 = paragraph.add_run(after)
    for r in (r0, r1, r2):
        r.font.name = "Times New Roman"
        r.font.size = Pt(11.0)
    r1.font.underline = WD_UNDERLINE.SINGLE


def _apply_ips_proposal_quote_formatting(doc: Document) -> None:
    """
    Restyle the standard estimate proposal body to match the IPS quote letterhead.

    Paragraph indices follow ``assets/estimate_template_autofill_logo_updated.docx`` body order.
    """
    try:
        for i, p in enumerate(doc.paragraphs):
            raw = _paragraph_full_text(p) or ""
            t = raw.strip()
            if not t:
                continue

            is_banner_title = _is_quote_title_line(t) or (
                i == 1
                and "quote #:" not in t.lower()
                and len(t) > 6
                and t.lower().rstrip().endswith("quote")
            )
            if is_banner_title:
                _set_paragraph_text_merged(p, t)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _paragraph_set_shading(p, "0B4F8A")
                _runs_set_banner_white(p, size_pt=20.0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                continue

            if re.match(r"^Quote\s*#\s*:", t, flags=re.IGNORECASE):
                m = re.match(r"^(Quote\s*#\s*:)\s*(.*)$", t, flags=re.IGNORECASE)
                label = m.group(1) if m else "Quote #: "
                body = (m.group(2) if m else "").strip()
                _clear_paragraph_runs(p)
                r0 = p.add_run(label)
                r0.font.underline = WD_UNDERLINE.SINGLE
                if body:
                    p.add_run(" " + body)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _paragraph_set_shading(p, "0B4F8A")
                for r in p.runs:
                    r.font.name = "Arial"
                    r.font.size = Pt(14.0)
                    r.font.bold = True
                    r.font.italic = True
                    r.font.color.rgb = RGBColor(255, 255, 255)
                p.paragraph_format.space_after = Pt(4)
                continue

            if i == 3 and t.startswith("Customer:"):
                _paragraph_clear_shading(p)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _runs_set_font(p, name="Times New Roman", size_pt=11.0)
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(0)
                continue

            if i == 4 and ("Attn:" in t or "Contact:" in t or "Quote Amt" in t):
                if t.startswith("Attn:"):
                    _set_paragraph_text_merged(p, "Contact:" + t[5:])
                _paragraph_clear_shading(p)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                _runs_set_font(p, name="Times New Roman", size_pt=11.0)
                p.paragraph_format.space_after = Pt(6)
                _paragraph_set_bottom_border_double(p)
                continue

            if i == 5 and "Industrial Plant Solutions" in t and "proposes" in t:
                _paragraph_clear_shading(p)
                _format_intro_solutions_underline(p)
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(8)
                continue

            if i == 6:
                _paragraph_clear_shading(p)
                _runs_set_font(p, name="Times New Roman", size_pt=11.0)
                p.paragraph_format.space_after = Pt(4)
                continue

            if i == 14 and "Responsibilities" in t:
                _paragraph_clear_shading(p)
                _runs_set_font(p, name="Times New Roman", size_pt=12.0, bold=True, italic=True)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(4)
                continue

            if i == 15:
                _paragraph_clear_shading(p)
                _runs_set_font(p, name="Times New Roman", size_pt=11.0)
                p.paragraph_format.space_after = Pt(6)
                continue

            if i == 18 and t.startswith("Prepared By:"):
                _paragraph_clear_shading(p)
                _runs_set_font(p, name="Times New Roman", size_pt=11.0)
                p.paragraph_format.space_before = Pt(14)
                continue

            if i == 19 and re.match(r"^Date:\s*", t):
                _paragraph_clear_shading(p)
                _runs_set_font(p, name="Times New Roman", size_pt=11.0)
                p.paragraph_format.space_after = Pt(10)
                continue

            if i == 20 and t.startswith("If you have any questions"):
                _paragraph_clear_shading(p)
                _paragraph_set_top_border_double(p)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                _set_paragraph_text_merged(p, t)
                _runs_set_font(p, name="Times New Roman", size_pt=11.0, bold=True, italic=True)
                p.paragraph_format.space_before = Pt(8)
    except Exception:
        return


def _strip_legacy_customer_location_tokens(text: str) -> str:
    """Remove spaced variants of ``{{CUSTOMER_LOCATION}}`` left in older templates."""
    if not text:
        return text
    return re.sub(r"\{\{\s*CUSTOMER_LOCATION\s*\}\}", "", text, flags=re.IGNORECASE)


def _fill_proposal_docx_from_bytes(raw: bytes, vals: dict[str, str]) -> bytes:
    repl = _docx_placeholder_replacements(vals)
    doc = Document(BytesIO(raw))
    _normalize_docx_placeholders(doc)

    def _strip_tokens(paragraph_text: str) -> str:
        return _strip_legacy_customer_location_tokens(paragraph_text)

    _apply_to_paragraph_text(doc, _strip_tokens)
    _replace_placeholders_doc(doc, repl)
    _remove_customer_location_paragraphs(doc)
    _apply_ips_proposal_quote_formatting(doc)
    _apply_standard_proposal_branding(doc)
    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()


def _convert_docx_bytes_to_pdf(docx_bytes: bytes) -> bytes:
    """
    Convert filled .docx to PDF using LibreOffice/soffice (any OS) or docx2pdf (Windows + Word).
    """
    with tempfile.TemporaryDirectory() as td:
        tdir = Path(td)
        src = tdir / "proposal.docx"
        src.write_bytes(docx_bytes)
        pdf_out = tdir / "proposal.pdf"
        for exe in ("soffice", "libreoffice"):
            try:
                subprocess.run(
                    [exe, "--headless", "--convert-to", "pdf", "--outdir", str(tdir), str(src)],
                    check=True,
                    timeout=180,
                    capture_output=True,
                    text=True,
                )
            except (FileNotFoundError, subprocess.CalledProcessError, subprocess.TimeoutExpired, OSError):
                continue
            if pdf_out.is_file():
                return pdf_out.read_bytes()

        try:
            from docx2pdf import convert  # type: ignore

            convert(str(src.resolve()))
            win_pdf = src.with_suffix(".pdf")
            if win_pdf.is_file():
                return win_pdf.read_bytes()
        except Exception:
            pass

    raise RuntimeError(
        "Could not convert the proposal Word file to PDF. Install **LibreOffice** (command `soffice` or "
        "`libreoffice` on PATH), or on Windows install **Microsoft Word** and run `pip install docx2pdf`."
    )


def try_convert_proposal_docx_to_pdf(docx_bytes: bytes) -> tuple[bytes | None, str]:
    """
    Convert a filled proposal .docx to PDF without raising.

    Returns ``(pdf_bytes, "")`` on success, or ``(None, short_user_message)`` on failure.
    """
    try:
        return _convert_docx_bytes_to_pdf(docx_bytes), ""
    except RuntimeError as e:
        msg = str(e).strip() or PROPOSAL_PDF_UNAVAILABLE_MSG
        return None, msg
    except Exception as e:
        return None, f"{PROPOSAL_PDF_UNAVAILABLE_MSG} ({type(e).__name__}: {e})"


def try_build_proposal_pdf(
    est: dict,
    totals: dict,
    customer_name: str = "",
    job_name: str = "",
    *,
    contact_name: str = "",
    prepared_by_phone: str = "",
    docx_bytes: bytes | None = None,
) -> tuple[bytes | None, str]:
    """
    Build proposal PDF from the same filled DOCX pipeline, without raising on conversion failure.

    If ``docx_bytes`` is provided, skips rebuilding the DOCX (caller already built it).

    Returns ``(pdf_bytes, "")`` on success, or ``(None, error_message)`` if the Word file could not be
    built or PDF conversion failed. Word-only failures return a message about the template/build step.
    """
    if docx_bytes is None:
        try:
            docx_bytes = build_proposal_docx(
                est,
                totals,
                customer_name=customer_name,
                job_name=job_name,
                contact_name=contact_name,
                prepared_by_phone=prepared_by_phone,
            )
        except FileNotFoundError as e:
            return None, str(e)
        except Exception as e:
            return None, f"Could not build the Word proposal: {type(e).__name__}: {e}"

    pdf, err = try_convert_proposal_docx_to_pdf(docx_bytes)
    if pdf is None and not err:
        return None, PROPOSAL_PDF_UNAVAILABLE_MSG
    return pdf, err


def build_proposal_docx(
    est: dict,
    totals: dict,
    customer_name: str = "",
    job_name: str = "",
    *,
    contact_name: str = "",
    prepared_by_phone: str = "",
    placeholder_values: dict[str, str] | None = None,
) -> bytes:
    """
    Build the proposal .docx from **assets/estimate_template_autofill_logo_updated.docx** only
    (no template overrides), replace text placeholders, then apply standard **company_logo** /
    ``{{COMPANY_LOGO}}`` branding from ``assets/`` per :data:`COMPANY_LOGO_CANDIDATE_FILENAMES`.

    Pass ``placeholder_values`` from a single :func:`proposal_values` call when the caller already
    computed it (avoids duplicate work and guarantees preview/download use the same map).
    """
    if placeholder_values is not None:
        vals = placeholder_values
    else:
        vals = proposal_values(
            est,
            totals,
            customer_name=customer_name,
            job_name=job_name,
            contact_name=contact_name,
            prepared_by_phone=prepared_by_phone,
        )
    raw = _load_standard_proposal_template_bytes()
    return _fill_proposal_docx_from_bytes(raw, vals)


def build_proposal_pdf(
    est: dict,
    totals: dict,
    customer_name: str = "",
    job_name: str = "",
    *,
    contact_name: str = "",
    prepared_by_phone: str = "",
) -> bytes:
    """
    Build the proposal PDF from the **same** filled DOCX as downloads/preview (template + placeholders).

    Raises ``RuntimeError`` (or ``FileNotFoundError`` from the DOCX step) if PDF conversion is unavailable;
    for non-throwing behavior use :func:`try_build_proposal_pdf` or :func:`try_convert_proposal_docx_to_pdf`.
    """
    pdf, err = try_build_proposal_pdf(
        est,
        totals,
        customer_name=customer_name,
        job_name=job_name,
        contact_name=contact_name,
        prepared_by_phone=prepared_by_phone,
        docx_bytes=None,
    )
    if pdf is None:
        if err and "Proposal template not found" in err:
            raise FileNotFoundError(err)
        raise RuntimeError(err or PROPOSAL_PDF_UNAVAILABLE_MSG)
    return pdf
